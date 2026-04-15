"""
Export Service — SkillLens
===========================
Generates an ATS-optimised, single-column .docx resume from
the Auto Rewrite output (a single plain-text string).

Uses python-docx. The document is built entirely in memory (io.BytesIO)
so nothing is ever written to disk.

ATS Design Rules Applied:
  - Single column, no tables, no text boxes, no headers/footers
  - Standard headings (Heading 2 style) for section names
  - Calibri 11pt body, 12pt headings
  - 0.75" margins on all sides
  - No images, no special characters that break OCR
"""

import io
import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


# ── Heading detection ─────────────────────────────────────────────────────────

# Standard resume section headings (case-insensitive match)
_SECTION_HEADINGS = {
    "summary", "professional summary", "career summary", "objective",
    "career objective", "profile", "professional profile", "about me",
    "experience", "professional experience", "work experience",
    "employment history", "work history", "employment",
    "education", "academic background", "academic history",
    "skills", "technical skills", "core competencies", "key skills",
    "competencies", "technologies", "tools & technologies",
    "projects", "projects & open source", "projects & open source contributions",
    "personal projects", "side projects", "open source",
    "certifications", "certificates", "licenses & certifications",
    "training", "professional development",
    "awards", "honors", "achievements", "accomplishments",
    "publications", "research", "volunteer", "volunteering",
    "languages", "interests", "references",
    "activities", "leadership", "extracurricular",
    "key technical skills", "key technical skills (ats-optimised)",
}

_BULLET_CHARS = set("•●·-*–—▪▸►➤➢■□")


def _is_section_heading(line: str) -> bool:
    """Check if a line is a resume section heading."""
    stripped = line.strip()

    # Too long to be a heading
    if len(stripped) > 60:
        return False

    # Exact match (case-insensitive, strip trailing colons/dashes)
    normalised = re.sub(r'[\s:—\-_]+$', '', stripped).lower()
    if normalised in _SECTION_HEADINGS:
        return True

    # ALL-CAPS lines under 50 chars that contain letters are likely headings
    if stripped.isupper() and len(stripped) < 50 and any(c.isalpha() for c in stripped):
        normalised_upper = re.sub(r'[\s:—\-_]+$', '', stripped).lower()
        if normalised_upper in _SECTION_HEADINGS:
            return True
        # Short ALL-CAPS line with no bullet char → treat as heading
        if stripped[0] not in _BULLET_CHARS and len(stripped) < 40:
            return True

    return False


def _is_bullet_line(line: str) -> bool:
    """Check if a line starts with a bullet character."""
    stripped = line.strip()
    if not stripped:
        return False
    if stripped[0] in _BULLET_CHARS:
        return True
    # Numbered list: "1." or "1)"
    if re.match(r'^\d{1,2}[.)]\s', stripped):
        return True
    return False


def _strip_bullet_char(line: str) -> str:
    """Remove the leading bullet character and whitespace."""
    stripped = line.strip()
    if stripped and stripped[0] in _BULLET_CHARS:
        return stripped[1:].strip()
    # Numbered list
    m = re.match(r'^\d{1,2}[.)]\s*', stripped)
    if m:
        return stripped[m.end():].strip()
    return stripped


def _add_heading(doc, text: str):
    """Add a bold, dark-navy section heading with a subtle bottom border."""
    p = doc.add_paragraph()
    p.space_before = Pt(14)
    p.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.font.name  = "Calibri"
    run.font.size  = Pt(12)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)  # dark navy

    # Bottom border via XML (ATS-safe — no tables needed)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'):  '6',
        qn('w:space'): '1',
        qn('w:color'): '4A5568',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_bullet(doc, text: str):
    """Add a bullet point paragraph."""
    p = doc.add_paragraph(style='List Bullet')
    p.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)


def _add_body(doc, text: str):
    """Add a normal body paragraph."""
    p = doc.add_paragraph()
    p.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)


# ── Main generator ────────────────────────────────────────────────────────────

def generate_ats_docx(rewritten_text: str) -> io.BytesIO:
    """
    Build an ATS-optimised .docx resume from the Auto Rewrite output.

    The function intelligently parses the plain-text string:
      - Lines matching known resume headings → Heading style
      - Lines starting with bullet characters → List Bullet style
      - Everything else → Normal body paragraph

    Parameters
    ----------
    rewritten_text : str
        The complete, plain-text resume produced by Auto Rewrite.

    Returns
    -------
    io.BytesIO
        In-memory byte stream of the .docx file, seeked to 0.
    """
    doc = Document()

    # ── Set 0.75" margins ─────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    # ── Default font ──────────────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.line_spacing = Pt(14)

    # ── Parse lines ───────────────────────────────────────────────
    lines = rewritten_text.split('\n')

    for line in lines:
        stripped = line.strip()

        # Skip empty lines (preserve spacing naturally)
        if not stripped:
            continue

        # Heading?
        if _is_section_heading(stripped):
            _add_heading(doc, stripped)

        # Bullet?
        elif _is_bullet_line(stripped):
            clean = _strip_bullet_char(stripped)
            if clean:
                _add_bullet(doc, clean)

        # Normal text
        else:
            _add_body(doc, stripped)

    # ── Save to in-memory stream ──────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
